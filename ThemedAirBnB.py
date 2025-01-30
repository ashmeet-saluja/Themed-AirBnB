# Combined Script for Airbnb Analysis

# 1. Data Collection (Datacollection.py)
# =====================================

import praw
import datetime

# Reddit API Credentials
reddit = praw.Reddit(user_agent='----',
                     client_id='-----', client_secret="-----",
                     username='------', password='-----')

# Specify the subreddit
subreddit = "Movies"

# Open file to write collected data
f = open('output1.csv', 'w', encoding='utf8')
f.write("MsgID, Timestamp,Author,ThreadID,ThreadTitle,MsgBody,ReplyTo,Permalink\n")

# Stream comments from the subreddit
count = 1
for comment in reddit.subreddit(subreddit).stream.comments():
    commentID = str(comment.id)
    author = str(comment.author).replace(";", "").replace("'", "").replace(",", "").replace("\"", "").replace("\n", " ").replace("\r", " ")
    timestamp = str(datetime.datetime.fromtimestamp(comment.created))
    replyTo = str(comment.parent().id) if not comment.is_root else "-"
    threadID = str(comment.submission.id)
    threadTitle = str(comment.submission.title).replace(";", "").replace("'", "").replace(",", "").replace("\"", "").replace("\n", " ").replace("\r", " ")
    msgBody = str(comment.body).replace(";", "").replace("'", "").replace(",", "").replace("\"", "").replace("\n", " ").replace("\r", " ")
    permalink = str(comment.permalink).replace(";", "").replace("'", "").replace(",", "").replace("\"", "").replace("\n", " ").replace("\r", " ")

    # Print collected message data to console
    print("-------------------------------------------------------")
    print("Comment ID: " + str(comment.id))
    print("Comment Author: " + str(comment.author))
    print("Timestamp: " + str(datetime.datetime.fromtimestamp(comment.created)))
    if not comment.is_root:
        print("Comment is a reply to: " + str(comment.parent().id))
    else:
        print("Comment is a reply to: -")
    print("Comment Thread ID: " + str(comment.submission.id))
    print("Comment Thread Title: " + str(comment.submission.title))
    print("Comment Body: " + str(comment.body))
    print("Comment Permalink: " + str(comment.permalink))

    # Write everything to a file
    f.write(f"'{commentID}','{timestamp}','{author}','{threadID}','{threadTitle}','{msgBody}','{replyTo}','{permalink}'\n")
    print("Total messages collected from /r/" + subreddit + ": " + str(count))
    count += 1


# 2. Text Processing
# ==================

# Data analysis libraries
import pandas as pd
import matplotlib.pyplot as plt
import seaborn as sns

# Data Preprocessing and Feature Engineering
from textblob import TextBlob
import re
import nltk
nltk.download('wordnet')
nltk.download('punkt')
nltk.download('stopwords')
from nltk.corpus import stopwords
from nltk.stem.wordnet import WordNetLemmatizer
from sklearn.feature_extraction.text import CountVectorizer, TfidfTransformer

# Model Selection and Validation
from sklearn.naive_bayes import MultinomialNB
from sklearn.model_selection import train_test_split
from sklearn.pipeline import Pipeline
from sklearn.metrics import confusion_matrix, classification_report, accuracy_score

# Import data
input_csv = pd.read_csv('/Users/nishkarshgupta/Desktop/ASU Coursework/Q4/Big Data/Files/CombinedMatt.csv', encoding='ISO-8859-1')

# Parse csv's to only extract column called 'message'
messages = input_csv['MsgBody']

# Text pre-processing functions
def text_processing(message):
    # Generating the list of words in the message (hashtags and other punctuations removed) and convert to lowercase
    def form_sentence(message):
        message = message.lower()  # Make messages lowercase
        message_blob = TextBlob(message.lower())  # Convert to 'textblob' which provides a simple API for NLP tasks
        return ' '.join(message_blob.words)
    
    new_message = form_sentence(message)
    
    # Removing stopwords and words with unusual symbols
    def no_user_alpha(message):
        message_list = [item for item in message.split()] 
        clean_words = [word for word in message_list if re.match(r'[^\W\d]*$', word)]  # remove punctuation and strange characters
        clean_sentence = ' '.join(clean_words) 
        clean_mess = [stopword for stopword in clean_sentence.split() if stopword not in stopwords.words('english')]  # remove stopwords
        return clean_mess
    
    no_punc_message = no_user_alpha(new_message)
    
    # Normalizing the words in messages 
    def normalization(message_list):
        lem = WordNetLemmatizer()
        normalized_message = []
        for word in message_list:
            normalized_text = lem.lemmatize(word, 'v')  # lemmatize words
            normalized_message.append(normalized_text)
        return normalized_message
    
    return normalization(no_punc_message)

# Print to console and write to file
f = open('processed_text_CM.csv', 'w', encoding='utf8')
for message in messages: 
    message = text_processing(message)
    for term in message:
        f.write(term + " ")
    f.write("\n")


# 3. LDA (LDA_Example.py)
# =======================

import pandas as pd  # Provides text processing capabilities
import numpy as np  # Provides Python with better math processing capabilities
from sklearn.feature_extraction.text import CountVectorizer

# The next line of code reads your Reddit data into this program's memory
# Place your reddit data into the same directory of this script and change the below filename
reviews_datasets = pd.read_csv('/Users/nishkarshgupta/Desktop/ASU Coursework/Q4/Big Data/processed_text.csv')

reviews_datasets = reviews_datasets.head(20000)  # Adjust this number according to the size of your dataset and whether you run into memory limitations
reviews_datasets.dropna()  # Drops any records that have a missing value

reviews_datasets.head()  # Print first 5 rows to console inspect data 

# This specifies which column to extract for text analysis
reviews_datasets['MsgBody'][10]

count_vect = CountVectorizer(max_df=0.8, min_df=2, stop_words='english')  # Hyperparameters
doc_term_matrix = count_vect.fit_transform(reviews_datasets['MsgBody'].values.astype('U'))  # Create document-term matrix

from sklearn.decomposition import LatentDirichletAllocation  # Import LDA

# n_components is how many topics you want to generate
LDA = LatentDirichletAllocation(n_components=10, random_state=42)  # n_components = number of topics to generate
LDA.fit(doc_term_matrix)

first_topic = LDA.components_[0]
top_topic_words = first_topic.argsort()[-10:]

# Prints out the most "important" words for forming topic distribution
print("Most \"Important\" words for forming topic distribution")  
for i in top_topic_words:
    print(count_vect.get_feature_names_out()[i])

for i, topic in enumerate(LDA.components_):
    print(f'Top 10 words for topic #{i}:')
    print([count_vect.get_feature_names_out()[i] for i in topic.argsort()[-10:]])
    print('\n')
    
topic_values = LDA.transform(doc_term_matrix)
topic_values.shape


# 4. LDA using bigram/trigram (LDA bigram_trigram model.py)
# =========================================================

import nltk; nltk.download('stopwords')
import re
import numpy as np
import pandas as pd
from pprint import pprint
import gensim
import gensim.corpora as corpora
from gensim.utils import simple_preprocess
from gensim.models import CoherenceModel
import spacy
import pyLDAvis
import pyLDAvis.gensim
import matplotlib.pyplot as plt

%matplotlib inline

import logging
logging.basicConfig(format='%(asctime)s : %(levelname)s : %(message)s', level=logging.ERROR)

import warnings
warnings.filterwarnings("ignore", category=DeprecationWarning)

from nltk.corpus import stopwords
stop_words = stopwords.words('english')
stop_words.extend(['from', 'subject', 're', 'edu', 'use'])

# Import Dataset
df = pd.read_csv('merged_reddit_data.csv', encoding='ISO-8859-1')

# Convert to list
data = df.MsgBody.values.tolist()

# Remove new line characters
data = [re.sub('\s+', ' ', sent) for sent in data]

# Remove distracting single quotes
data = [re.sub("\'", "", sent) for sent in data]

def sent_to_words(sentences):
    for sentence in sentences:
        yield(gensim.utils.simple_preprocess(str(sentence), deacc=True))  # deacc=True removes punctuations

data_words = list(sent_to_words(data))

# Build the bigram and trigram models
bigram = gensim.models.Phrases(data_words, min_count=3, threshold=100)
trigram = gensim.models.Phrases(bigram[data_words], threshold=100)

bigram_mod = gensim.models.phrases.Phraser(bigram)
trigram_mod = gensim.models.phrases.Phraser(trigram)

def remove_stopwords(texts):
    return [[word for word in simple_preprocess(str(doc)) if word not in stop_words] for doc in texts]

def make_bigrams(texts):
    return [bigram_mod[doc] for doc in texts]

def make_trigrams(texts):
    return [trigram_mod[bigram_mod[doc]] for doc in texts]

def lemmatization(texts, allowed_postags=['NOUN', 'ADJ', 'VERB', 'ADV']):
    texts_out = []
    for sent in texts:
        doc = nlp(" ".join(sent)) 
        texts_out.append([token.lemma_ for token in doc if token.pos_ in allowed_postags])
    return texts_out

data_words_nostops = remove_stopwords(data_words)
data_words_bigrams = make_bigrams(data_words_nostops)
nlp = spacy.load('en_core_web_sm', disable=['parser', 'ner'])
data_lemmatized = lemmatization(data_words_bigrams, allowed_postags=['NOUN', 'ADJ', 'VERB', 'ADV'])

id2word = corpora.Dictionary(data_lemmatized)
texts = data_lemmatized
corpus = [id2word.doc2bow(text) for text in texts]

print(corpus[:1])

lda_model = gensim.models.ldamodel.LdaModel(corpus=corpus,
                                            id2word=id2word,
                                            num_topics=10, 
                                            random_state=100,
                                            update_every=1,
                                            chunksize=100,
                                            passes=10,
                                            alpha='auto',
                                            per_word_topics=True)

pprint(lda_model.print_topics())
doc_lda = lda_model[corpus]


# 5. Seeded LDA (Seeded LDA_Movies.py)
# ====================================


from gensim import corpora
from gensim.models import LdaModel
from gensim.utils import simple_preprocess
from gensim.parsing.preprocessing import STOPWORDS
import numpy as np
import pandas as pd

# Read documents from CSV file
df = pd.read_csv('merged_reddit_data.csv')  # Replace 'your_file.csv' with the path to your CSV file
documents = df['MsgBody'].tolist()  # Assuming 'text_column' contains your documents

# Preprocess the documents
texts = [simple_preprocess(doc) for doc in documents]

# Define seed words for each topic
seed_words = {
    "Topic 0": ["Star Wars", "planets", "dark side", "system"],
    "Topic 1": ["Wizard Oz", "yellow", "brick", "good witch"],
    "Topic 2": ["Lord Rings", "wizard", "journey", "Frodo"],
    "Topic 3": ["ET", "phone", "home", "Spielberg", "Barrymore"],
    "Topic 4": ["Snow White", "dwarfs", "poison", "singing"],
    "Topic 5": ["Terminator", "skynet", "Judgement Day", "Schwarzenegger"],
    "Topic 6": ["Lion King", "Nala", "Rafiki", "Mufasa", "musical"],
    "Topic 7": ["Godfather", "gun", "canoli", "gangster", "organized crime"],
    "Topic 8": ["Jesus Film", "Israel", "bible", "Christ", "crucifi"],
    "Topic 9": ["Jurassic Park", "Dinosaur", "clone", "T-rex"],
}

# Create dictionary and corpus
dictionary = corpora.Dictionary(texts)
corpus = [dictionary.doc2bow(text) for text in texts]

# Initialize the LDA model without seed_topics
lda_model = LdaModel(
    corpus=corpus,
    id2word=dictionary,
    num_topics=len(seed_words),
    passes=10,
    random_state=42,
    iterations=1000  # Increase iterations for better convergence
)

# Manually set topic-word distributions using seed words
num_topics = lda_model.num_topics
num_words = len(dictionary)
topic_dist = np.zeros((num_topics, num_words))  # Initialize topic-word distributions matrix
for topic_name, words in seed_words.items():
    topic_id = int(topic_name.split()[1])  # Extract numerical part of topic name
    topic = [dictionary.token2id[word] for word in words if word in dictionary.token2id]
    topic_dist[topic_id, :] = 0.0  # Set all values for the topic to 0
    topic_dist[topic_id, topic] = 1.0  # Set the seed words to 1

# Update the model with the modified topic distributions
lda_model.expElogbeta[:] = topic_dist

# Train the model
lda_model.update(corpus)

# Print the topics
for idx, topic in lda_model.print_topics(-1):
    print('Topic: {} \nWords: {}'.format(idx, topic))


# 6. Clustering (Bert Clustering.py)
# ===================================

import pandas as pd
from sklearn.cluster import KMeans
from transformers import BertTokenizer, BertModel
import torch
import numpy as np

# Load data
reviews_datasets = pd.read_csv('processed_text.csv')
reviews_datasets.dropna(inplace=True)

# Initialize BERT tokenizer and model
tokenizer = BertTokenizer.from_pretrained('bert-base-uncased')
model = BertModel.from_pretrained('bert-base-uncased')

# Function to encode text into BERT embeddings
def get_bert_embeddings(texts):
    model.eval()
    with torch.no_grad():
        inputs = tokenizer(texts, return_tensors="pt", padding=True, truncation=True, max_length=512)
        outputs = model(**inputs)
        # Take the first token (CLS token) embeddings from each sequence
        embeddings = outputs.last_hidden_state[:, 0, :].numpy()
    return embeddings

# Encode all reviews
embeddings = np.vstack([get_bert_embeddings(text) for text in reviews_datasets['MsgBody']])

# Clustering with K-means
kmeans = KMeans(n_clusters=10, random_state=42)
kmeans.fit(embeddings)
clusters = kmeans.labels_

# Print the clusters
for i in range(10):
    cluster_texts = reviews_datasets['MsgBody'][clusters == i]
    print(f"Cluster {i} examples:")
    print(cluster_texts.head())
    print('\n')


# 7. Sentiment Analysis with BERT/Roberta (Roberta_twitter.py)
# ============================================================


import torch
from transformers import RobertaTokenizer, RobertaForSequenceClassification

def roberta_sentiment_analysis(text):
    # Load pre-trained RoBERTa tokenizer and model
    tokenizer = RobertaTokenizer.from_pretrained('cardiffnlp/twitter-roberta-base-sentiment-latest')
    model = RobertaForSequenceClassification.from_pretrained('cardiffnlp/twitter-roberta-base-sentiment-latest', num_labels=3)  # Assuming binary classification (positive/negative)

    # Tokenize input text
    inputs = tokenizer(text, return_tensors='pt', padding=True, truncation=True)

    # Forward pass through the model
    outputs = model(**inputs)

    # Get predicted probabilities and predicted class
    probabilities = torch.softmax(outputs.logits, dim=1)
    predicted_class = torch.argmax(probabilities, dim=1).item()

    # Decode predicted class
    if predicted_class == 0:
        sentiment = 'Negative'
    else:
        sentiment = 'Positive'

    return sentiment, probabilities[0][predicted_class].item()

# Read comments from the text file
comments_file_path = 'comments_with_keyword.txt'  # Change this to the path of your text file
with open(comments_file_path, 'r', encoding='utf-8') as file:
    comments = file.readlines()
comments = [comment.strip() for comment in comments]

# Initialize counters for positive and negative sentiments
positive_count = 0
negative_count = 0

# Iterate through each comment
for comment in comments:
    # Apply RoBERTa sentiment analysis to the comment
    sentiment, confidence = roberta_sentiment_analysis(comment)
    
    # Aggregate sentiment results
    if sentiment == 'Positive':
        positive_count += 1
    elif sentiment == 'Negative':
        negative_count += 1

# Calculate overall sentiment
total_comments = len(comments)
positive_percentage = (positive_count / total_comments) * 100
negative_percentage = (negative_count / total_comments) * 100

print("Sentiment analysis results for Reddit comments related to 'prometheus':")
print("Total comments:", total_comments)
print("Positive comments:", positive_count, f"({positive_percentage:.2f}%)")
print("Negative comments:", negative_count, f"({negative_percentage:.2f}%)")


# 8. Keyword Extraction for Roberta (Key word pull for Roberta.py)
# ================================================================


import csv

def search_comments(csv_file, keyword, output_file):
    found_comments = []

    with open(csv_file, 'r', encoding='utf-8') as file:
        reader = csv.DictReader(file)
        for row in reader:
            comment = row['MsgBody']  # Assuming the column name is 'Comment'
            # Check if the comment contains the keyword
            if keyword.lower() in comment.lower():
                # Remove single quotes from the comment and split it into sentences
                sentences = comment.strip("'").split('.')
                # Format each sentence and add it to the list of found comments
                for sentence in sentences:
                    formatted_sentence = f'"{sentence.strip()}",'
                    found_comments.append(formatted_sentence)

    # Format the list of comments containing the keyword as a string with each sentence on a new line
    formatted_comments = '[' + '\n'.join(found_comments) + ']'

    # Write the formatted string to the output text file
    with open(output_file, 'w', encoding='utf-8') as output:
        output.write(formatted_comments)

# Example usage
csv_file = 'input.csv'  # Replace 'input.csv' with the path to your CSV file
keyword = 'prometheus'  # Replace 'alien' with the keyword from your LDA model results
output_file = 'comments_with_keyword.txt'  # Specify the output file name

search_comments(csv_file, keyword, output_file)
print(f"Comments containing the keyword '{keyword}' have been saved to '{output_file}'.")


# 9. Final Analysis using LangChain with Ollama2 (OllamaCount.py)
# ===============================================================

import pandas as pd
from docx import Document
from langchain_community.llms import Ollama
from collections import Counter

# Initialize the Llama2 model on Ollama
llm = Ollama(base_url='http://localhost:11434', model="llama2")

# Load your CSV data
reddit_data = pd.read_csv('/Users/nishkarshgupta/Desktop/ASU Coursework/Q4/Big Data/Files/CombinedMatt.csv')

# Combine ThreadTitle and MsgBody if both exist, otherwise use what's available
reddit_data['text'] = reddit_data.apply(
    lambda row: row['ThreadTitle'] + " " + row['MsgBody'] if pd.notnull(row['MsgBody']) else row['ThreadTitle'], axis=1)

# Create a new Word document for results
results_doc = Document()

# Function to classify text and add results to the Word document
def classify_text(text):
    response = llm.invoke(f"Classify the following text into categories: television, movies, books. Text: {text}")
    if response:
        results_doc.add_paragraph(f"Text: {text}\nClassification: {response.strip()}\n")
        return response.strip()
    else:
        error_msg = "Error: Unable to classify text"
        results_doc.add_paragraph(f"Text: {text}\nClassification: {error_msg}\n")
        return error_msg

# Classify texts and add to document for the first 100 entries of the DataFrame using slicing
for text in reddit_data['text'][0:100]:
    classify_text(text)

# Save the document with classification results
results_doc.save('/Users/nishkarshgupta/Desktop/ASU Coursework/Q4/Big Data/Files/ClassificationResults_300.docx')

# Load the document to process counts
doc = Document('/Users/nishkarshgupta/Desktop/ASU Coursework/Q4/Big Data/Files/ClassificationResults_300.docx')

# Initialize counters for each category
television_shows = Counter()
movies = Counter()
books = Counter()

# Define functions to add to the counters
def add_to_counter(counter, item):
    if item:
        counter[item] += 1

# Process the document to fill counters
for para in doc.paragraphs:
    parts = para.text.split('*')
    for part in parts:
        if 'Television:' in part:
            show_name = part.split('"')[1] if '"' in part else None
            add_to_counter(television_shows, show_name)
        elif 'Movies:' in part:
            movie_name = part.split('"')[1] if '"' in part else None
            add_to_counter(movies, movie_name)
        elif 'Books:' in part:
            book_name = part.split('"')[1] if '"' in part else None
            add_to_counter(books, book_name)

# Create another document for counts
counts_doc = Document()

# Function to write results to document
def write_results_to_doc(counter, title, document):
    document.add_heading(title, level=1)
    for item, count in counter.items():
        document.add_paragraph(f'{item} - {count}')

# Write simplified results to the document
write_results_to_doc(television_shows, "Television Shows", counts_doc)
write_results_to_doc(movies, "Movies", counts_doc)
write_results_to_doc(books, "Books", counts_doc)

# Save the document with counts
counts_doc.save('/Users/nishkarshgupta/Desktop/ASU Coursework/Q4/Big Data/Files/CountsResults_300.docx')

# NOTE: Replace 'your_client_id', 'your_client_secret', 'your_user_agent', 'your_api_key' with actual values.
# Make sure to add the actual logic for vectorization, model loading, and other placeholders used in this script.
